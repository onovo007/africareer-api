"""
AfriCareer AI - core engine (framework-agnostic).

All business logic lives here: OpenAI + Pinecone + Tavily, RAG retrieval, CV /
cover-letter / motivation-letter generation, verified course/job/opportunity search.

No Streamlit, no FastAPI - so it can be imported by the FastAPI service (api.py),
the existing Streamlit app, tests, or a future worker. Configure via env vars:
OPENAI_API_KEY, PINECONE_API_KEY, TAVILY_API_KEY (optional).
"""
import os
import io
import json
import re
from datetime import datetime
from functools import lru_cache
from urllib.parse import quote_plus, urlparse

import httpx
from dotenv import load_dotenv
load_dotenv()

from pinecone import Pinecone, ServerlessSpec
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

APP_NAME = "AfriCareer AI"
INDEX_NAME = "africareer-kb"

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()

# ---------------------------------------------------------------- clients (lazy)
_pc = None
_index = None
_embeddings = None
_llm = None


def _clients():
    """Lazily create and cache the Pinecone index, embeddings, and LLM clients."""
    global _pc, _index, _embeddings, _llm
    if _index is None:
        _pc = Pinecone(api_key=PINECONE_API_KEY)
        if INDEX_NAME not in [i.name for i in _pc.list_indexes()]:
            _pc.create_index(name=INDEX_NAME, dimension=1536, metric="cosine",
                             spec=ServerlessSpec(cloud="aws", region="us-east-1"))
        _index = _pc.Index(INDEX_NAME)
    if _embeddings is None:
        _embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    if _llm is None:
        _llm = ChatOpenAI(temperature=0.7, model="gpt-4o-mini", openai_api_key=OPENAI_API_KEY)
    return _index, _embeddings, _llm


SAFETY_SYSTEM_MESSAGE = """You are AfriCareer AI, a comprehensive career and academic guidance assistant for African youth and professionals.

YOUR MISSION: Help with ANY career, job, education, scholarship, or professional development question. Be as helpful as possible.

You should answer questions about: career guidance and planning; CV/resume creation and improvement; job search and interviews; salary and offers; education pathways, courses and certifications; scholarships, university admissions and motivation letters; entrepreneurship; workplace issues; professional development; industry-specific guidance; and questions about AfDB, UNICEF, ILO and UNESCO frameworks.

You MUST refuse ONLY: sexual or explicit content; violence or instructions for illegal activities. For those, reply: "I'm AfriCareer AI for career and academic guidance. I can't help with that specific topic, but I'm here to help with any career, job, education, or scholarship question."

RESPONSE STYLE: be practical, specific and actionable; cite AfDB SEPA, UNICEF, ILO or UNESCO frameworks when relevant; ground advice in the African context; be supportive and encouraging; never invent facts about a person, employer, or institution."""


def safe_llm_call(user_prompt, rag_context="", language="English"):
    """Single LLM entry point with the safety system message and optional RAG grounding."""
    _, _, llm = _clients()
    system_message = SystemMessage(content=SAFETY_SYSTEM_MESSAGE)
    if rag_context:
        full_prompt = (f"Language: {language}\n\n"
                       f"CONTEXT FROM AUTHORITATIVE SOURCES (AfDB SEPA, UNICEF, ILO, UNESCO):\n{rag_context}\n\n"
                       f"USER REQUEST:\n{user_prompt}\n\n"
                       f"Provide your response in {language}, grounded in the context above. "
                       f'Cite specific frameworks when relevant (e.g., "According to AfDB SEPA...").')
    else:
        full_prompt = f"Language: {language}\n\nUSER REQUEST:\n{user_prompt}\n\nProvide your response in {language}."
    try:
        return llm.invoke([system_message, HumanMessage(content=full_prompt)]).content
    except Exception as e:
        return f"Error: {str(e)}"


def retrieve_career_guidance(query, top_k=5):
    """Retrieve grounding context from the Pinecone knowledge base."""
    index, embeddings, _ = _clients()
    try:
        query_vec = embeddings.embed_query(query)
        results = index.query(vector=query_vec, top_k=top_k, include_metadata=True)
        pieces, sources = [], []
        for match in results["matches"]:
            if match.get("metadata") and match.get("score", 0) > 0.7:
                pieces.append(match["metadata"]["text"])
                if "source" in match["metadata"]:
                    sources.append(match["metadata"]["source"])
        context = "\n\n".join(pieces) if pieces else ""
        if sources:
            context += f"\n\n[Sources: {', '.join(sorted(set(sources)))}]"
        return context
    except Exception:
        return ""


# ------------------------------------------------------------ verified web links
_BROWSER_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
               "(KHTML, like Gecko) Chrome/124.0 Safari/537.36")

_PROVIDER_SEARCH = {
    "coursera": "https://www.coursera.org/search?query={q}",
    "edx": "https://www.edx.org/search?q={q}",
    "udemy": "https://www.udemy.com/courses/search/?q={q}",
    "udacity": "https://www.udacity.com/catalog?searchValue={q}",
    "class central": "https://www.classcentral.com/search?q={q}",
    "classcentral": "https://www.classcentral.com/search?q={q}",
    "freecodecamp": "https://www.freecodecamp.org/news/search/?query={q}",
    "khan academy": "https://www.khanacademy.org/search?page_search_query={q}",
    "khanacademy": "https://www.khanacademy.org/search?page_search_query={q}",
    "linkedin learning": "https://www.linkedin.com/learning/search?keywords={q}",
    "youtube": "https://www.youtube.com/results?search_query={q}",
    "alison": "https://alison.com/courses?query={q}",
    "futurelearn": "https://www.futurelearn.com/search?q={q}",
}
_PAID_PROVIDERS = ("udemy", "udacity", "linkedin learning")
_FREE_PROVIDERS = ("freecodecamp", "khan academy", "khanacademy", "youtube",
                   "mit opencourseware", "ocw", "alison")
_JOB_BOARDS = ("linkedin.com", "indeed.com", "glassdoor.com", "ziprecruiter.com",
               "jobberman.com", "myjobmag.com", "brightermonday.co.ke", "careers24.com")
_NGO_IO_BOARDS = ("who.int", "unicef.org", "gavi.org", "unv.org", "un.org", "undp.org",
                  "unhcr.org", "worldbank.org", "fhi360.org", "path.org",
                  "reliefweb.int", "unjobs.org", "impactpool.org", "devex.com", "idealist.org")


def provider_search_url(provider, query):
    q = quote_plus((query or "").strip())
    p = (provider or "").lower()
    for key, tmpl in _PROVIDER_SEARCH.items():
        if key in p:
            return tmpl.format(q=q)
    return f"https://www.classcentral.com/search?q={q}"


def classify_cost(provider, llm_cost):
    p = (provider or "").lower()
    if any(k in p for k in _PAID_PROVIDERS):
        return "Paid"
    if any(k in p for k in _FREE_PROVIDERS):
        return "Free"
    lc = (llm_cost or "").lower()
    if "paid" in lc and "free" not in lc:
        return "Paid"
    return "Free"


def cost_matches(cost, pref):
    if pref.startswith("Free &"):
        return True
    if pref.startswith("Free"):
        return cost == "Free"
    if pref.startswith("Paid"):
        return cost == "Paid"
    return True


@lru_cache(maxsize=1024)
def verify_url(url, timeout=6.0):
    """True if reachable for a real user (2xx/3xx, or anti-bot 401/403/405/429/999)."""
    if not url or not url.startswith(("http://", "https://")):
        return False
    try:
        with httpx.Client(follow_redirects=True, timeout=timeout,
                          headers={"User-Agent": _BROWSER_UA}) as c:
            code = c.get(url).status_code
        return code < 400 or code in (401, 403, 405, 429, 999)
    except Exception:
        return False


def web_search_links(query, max_results=4):
    if not TAVILY_API_KEY:
        return []
    try:
        with httpx.Client(timeout=12.0) as c:
            resp = c.post("https://api.tavily.com/search",
                          json={"api_key": TAVILY_API_KEY, "query": query,
                                "max_results": max_results, "search_depth": "basic"})
        return [{"title": r.get("title", r.get("url", "")), "url": r.get("url", "")}
                for r in resp.json().get("results", []) if r.get("url")]
    except Exception:
        return []


def web_research(query, max_results=5):
    if not TAVILY_API_KEY:
        return ""
    try:
        with httpx.Client(timeout=15.0) as c:
            resp = c.post("https://api.tavily.com/search",
                          json={"api_key": TAVILY_API_KEY, "query": query,
                                "max_results": max_results, "search_depth": "advanced",
                                "include_answer": True})
            data = resp.json()
        parts = []
        if data.get("answer"):
            parts.append("Summary: " + data["answer"])
        for r in data.get("results", []):
            content = (r.get("content") or "").strip()
            if content:
                parts.append(f"- {r.get('title', '')}: {content[:400]}")
        return "\n".join(parts)[:4000]
    except Exception:
        return ""


def web_job_search(query, time_range="", domains=None, max_results=10):
    if not TAVILY_API_KEY:
        return []
    payload = {"api_key": TAVILY_API_KEY, "query": query,
               "max_results": max_results, "search_depth": "basic"}
    if time_range:
        payload["time_range"] = time_range
    if domains:
        payload["include_domains"] = list(domains)
    try:
        with httpx.Client(timeout=15.0) as c:
            resp = c.post("https://api.tavily.com/search", json=payload)
        return [{"title": r.get("title", r.get("url", "")), "url": r.get("url", ""),
                 "content": (r.get("content") or "")[:200]}
                for r in resp.json().get("results", []) if r.get("url")]
    except Exception:
        return []


def _job_label(r):
    title = " ".join((r.get("title") or "").split()).strip()
    domain = urlparse(r.get("url", "")).netloc.replace("www.", "")
    if len(title) < 3:
        title = f"Job posting on {domain or 'the web'}"
    return title, domain


def _extract_json(text):
    """Parse a JSON object or array from an LLM response, tolerating code fences/prose."""
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r'(\{[\s\S]*\}|\[[\s\S]*\])', text)
        if m:
            return json.loads(m.group(1))
        raise


# ------------------------------------------------------------------ DOCX writers
def generate_premium_cv_docx(cv_json_str):
    """Premium 2-page ATS-optimized CV as DOCX (bytes) from LLM JSON."""
    cv = _extract_json(cv_json_str)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    NAVY = RGBColor(0x0F, 0x2B, 0x4C)
    GRAY = RGBColor(0x4A, 0x55, 0x68)
    DARK = RGBColor(0x2D, 0x2D, 0x2D)

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8'); bottom.set(qn('w:color'), '14B8A6')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    def set_run(run, size=11, color=DARK, bold=False, italic=False, font_name="Georgia"):
        run.font.size = Pt(size); run.font.color.rgb = color
        run.bold = bold; run.italic = italic; run.font.name = font_name

    def heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(title.upper()), size=12, color=NAVY, bold=True)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:color'), '14B8A6')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    def bullet(text, size=10):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Cm(0.8)
        p.clear(); set_run(p.add_run(text), size=size, color=DARK)

    name_p = doc.add_paragraph(); name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(cv.get("full_name", "CANDIDATE NAME").upper())
    set_run(r, size=18, color=NAVY, bold=True); r.font.character_spacing = Pt(2)

    for key, size in (("credentials", 9.5), ("contact_line", 9.5)):
        val = cv.get(key, "")
        if val:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(val), size=size, color=GRAY, italic=(key == "credentials"))
    add_divider()

    if cv.get("professional_summary"):
        heading("Professional Summary")
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        set_run(p.add_run(cv["professional_summary"]), size=10.5, color=DARK)

    if cv.get("selected_achievements"):
        heading("Selected Achievements")
        for ach in cv["selected_achievements"]:
            bullet(ach)

    if cv.get("core_competencies"):
        heading("Core Competencies")
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        set_run(p.add_run("  •  ".join(cv["core_competencies"])), size=10, color=DARK)

    if cv.get("work_experience"):
        heading("Professional Experience")
        for job in cv["work_experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
            set_run(p.add_run(job.get("title", "")), size=11, color=NAVY, bold=True)
            if job.get("company"):
                set_run(p.add_run(f" - {job['company']}"), size=11, color=DARK)
            loc_date = [x for x in (job.get("location"), job.get("dates")) if x]
            if loc_date:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
                set_run(p.add_run(" | ".join(loc_date)), size=9.5, color=GRAY, italic=True)
            for b in job.get("bullets", []):
                bullet(b)

    if cv.get("education"):
        heading("Education")
        for edu in cv["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
            set_run(p.add_run(edu.get("degree", "")), size=11, color=NAVY, bold=True)
            if edu.get("institution"):
                p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(1)
                set_run(p2.add_run(edu["institution"]), size=10, color=DARK)
                if edu.get("dates"):
                    set_run(p2.add_run(f" - {edu['dates']}"), size=10, color=GRAY, italic=True)

    if cv.get("publications"):
        heading("Selected Publications")
        for pub in cv["publications"]:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(pub), size=10, color=DARK)

    if cv.get("projects"):
        heading("Selected Projects & Deployments")
        for proj in cv["projects"]:
            bullet(proj)

    if cv.get("certifications"):
        heading("Certifications & Training")
        for cert in cv["certifications"]:
            bullet(cert)

    if cv.get("technical_skills"):
        heading("Technical Skills")
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(cv["technical_skills"]), size=10, color=DARK)

    if cv.get("languages"):
        heading("Languages")
        p = doc.add_paragraph()
        set_run(p.add_run("  •  ".join(cv["languages"])), size=10, color=DARK)

    add_divider()
    footer = doc.add_paragraph(); footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run(f"Generated by {APP_NAME} • {datetime.now().strftime('%B %d, %Y')}"),
            size=8, color=GRAY, italic=True)

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()


def generate_premium_cover_letter_docx(letter_json_str):
    """Premium cover / motivation letter as DOCX (bytes) from LLM JSON."""
    cl = _extract_json(letter_json_str)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    NAVY = RGBColor(0x0F, 0x2B, 0x4C)
    GRAY = RGBColor(0x4A, 0x55, 0x68)
    DARK = RGBColor(0x2D, 0x2D, 0x2D)

    def set_run(run, size=11, color=DARK, bold=False, italic=False):
        run.font.size = Pt(size); run.font.color.rgb = color
        run.bold = bold; run.italic = italic; run.font.name = "Georgia"

    def teal_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(16)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8'); bottom.set(qn('w:color'), '14B8A6')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    name_p = doc.add_paragraph(); name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(cl.get("full_name", "CANDIDATE NAME").upper())
    set_run(r, size=16, color=NAVY, bold=True); r.font.character_spacing = Pt(2)

    for key in ("credentials", "contact_line"):
        val = cl.get(key, "")
        if val:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(val), size=9.5, color=GRAY, italic=(key == "credentials"))
    teal_divider()

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(cl.get("date", datetime.now().strftime("%B %d, %Y"))), size=11, color=DARK)

    for line in cl.get("addressee_lines", ["Hiring Committee"]):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(line), size=11, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if cl.get("re_line"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
        set_run(p.add_run("RE: "), size=11, color=NAVY, bold=True)
        set_run(p.add_run(cl["re_line"]), size=11, color=NAVY, bold=True)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(cl.get("salutation", "Dear Hiring Manager,")), size=11, color=DARK)

    for para in cl.get("body_paragraphs", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = Pt(15)
        set_run(p.add_run(para), size=11, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(cl.get("closing_line", "Respectfully submitted,")), size=11, color=DARK)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(cl.get("signature_name", cl.get("full_name", ""))), size=11, color=NAVY, bold=True)
    for key in ("signature_title", "signature_contact"):
        if cl.get(key):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(cl[key]), size=10, color=GRAY)

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()


# ---------------------------------------------------------- high-level operations
def career_guidance(answers, language="English"):
    """Career roadmap text from the 5-question profile."""
    ctx = retrieve_career_guidance(f"career paths employment skills development Africa {answers[:100]}")
    prompt = (f"Provide career guidance for African youth based on their profile.\n\n"
              f"Their Answers: {answers}\n\n"
              "Provide:\n1. Top 3 Career Paths in Africa (with reasons)\n"
              "2. Key Skills to Develop (7-10 with brief explanations)\n"
              "3. Action Plan (5 concrete steps)\n\n"
              "Ground advice in African job-market realities and cite frameworks when applicable.")
    return safe_llm_call(prompt, ctx, language)


def assistant_answer(question, language="English"):
    return safe_llm_call(question, retrieve_career_guidance(question), language)


def analyze_resume(resume_text, city="", additional_info="", language="English"):
    ctx = retrieve_career_guidance(f"resume improvement African job market {city or 'Africa'}")
    prompt = (f"Analyze this resume for the African job market.\n\nResume Content:\n{resume_text[:3000]}\n\n"
              f"Location Context: {city or 'General African market'}\nAdditional Info: {additional_info or 'None'}\n\n"
              "Provide:\n1. ATS Compatibility Score (1-100)\n2. Top 3 Strengths\n"
              "3. Top 5 Areas for Improvement\n4. African Market Relevance Assessment\n5. 3 Actionable Next Steps")
    return safe_llm_call(prompt, ctx, language)


_CV_SCHEMA = """{
  "full_name": "", "credentials": "", "contact_line": "",
  "professional_summary": "3-4 sentence summary, tailored to the African market",
  "selected_achievements": ["4-6 quantified achievements with numbers/scale/outcomes"],
  "core_competencies": ["8-12 ATS keyword-rich skills"],
  "work_experience": [{"title": "", "company": "", "location": "", "dates": "", "bullets": ["achievement bullets"]}],
  "education": [{"degree": "", "institution": "", "dates": ""}],
  "publications": [], "projects": [], "certifications": [],
  "technical_skills": "comma-separated tools", "languages": ["English (Native)"]
}"""


def build_cv_from_resume(resume_text, feedback="", language="English"):
    """Return DOCX bytes for an improved CV built from an existing resume + analysis feedback."""
    ctx = retrieve_career_guidance("professional CV resume best practices African job market ATS optimization")
    prompt = (f"You are a professional CV writer for the African job market.\n\n"
              f"ORIGINAL RESUME:\n{resume_text[:4000]}\n\nANALYSIS FEEDBACK:\n{feedback[:2000]}\n\n"
              f"Create an improved, ATS-optimized CV. RESPOND ONLY WITH VALID JSON using this structure:\n{_CV_SCHEMA}\n\n"
              "RULES: keep ALL factual details from the resume; do NOT invent experience; "
              "improve bullets with action verbs and metrics; never use placeholder brackets; return ONLY the JSON.")
    return generate_premium_cv_docx(safe_llm_call(prompt, ctx, "English"))


def build_cv_from_answers(answers, full_name="", contact_line="", language="English"):
    """Return DOCX bytes for a CV built from the 5-question profile (no invention)."""
    ctx = retrieve_career_guidance("professional CV resume best practices African job market ATS optimization")
    prompt = (f"You are a professional CV writer creating an ATS CV from a jobseeker's answers.\n\n"
              f"CANDIDATE ANSWERS:\n{answers}\n\n"
              f"CONTACT (use verbatim; do not invent): Full name: {full_name or '(not provided)'}; "
              f"Contact line: {contact_line or '(not provided)'}\n\n"
              f"RESPOND ONLY WITH VALID JSON using this structure:\n{_CV_SCHEMA}\n\n"
              "RULES: use ONLY facts the candidate provided; do NOT invent employers, dates, metrics, or degrees; "
              "NEVER output placeholder brackets like [Your Name] - omit unknown fields; return ONLY the JSON.")
    return generate_premium_cv_docx(safe_llm_call(prompt, ctx, "English"))


def build_cover_letter(resume_text, position, company, city=""):
    """Return DOCX bytes for a researched cover letter."""
    ctx = retrieve_career_guidance(f"cover letter professional {position} {company} African job market")
    org_research = web_research(f"{company} organization mission, products, values, and recent work") if TAVILY_API_KEY else ""
    prompt = (f"You are a professional cover letter writer.\n\nCANDIDATE'S RESUME:\n{resume_text[:4000]}\n\n"
              f"TARGET POSITION: {position}\nTARGET COMPANY: {company}\nLOCATION: {city or 'Africa'}\n\n"
              f"ORGANIZATION RESEARCH (verified; do not invent beyond this):\n{org_research or '(none available)'}\n\n"
              "RESPOND ONLY WITH VALID JSON: {"
              '"full_name": "", "credentials": "", "contact_line": "", '
              f'"date": "{datetime.now().strftime("%B %d, %Y")}", '
              f'"addressee_lines": ["Hiring Committee", "{company}"], "re_line": "{position} Position", '
              '"salutation": "Dear Hiring Manager,", "body_paragraphs": ['
              '"Opening: name the role and connect it to the org mission/priority from the research; one-sentence positioning.",'
              '"Map your most relevant experience to the role, with metrics from the resume.",'
              '"A second capability the role needs, with concrete evidence/validation.",'
              '"Specific knowledge of the organization (from research) and why you fit.",'
              '"Close: reaffirm interest, availability, invite an interview."], '
              '"closing_line": "Respectfully submitted,", "signature_name": "", "signature_title": "", "signature_contact": ""}\n\n'
              "RULES: use ONLY factual details from the resume; do NOT invent; return ONLY the JSON.")
    return generate_premium_cover_letter_docx(safe_llm_call(prompt, ctx, "English"))


def build_motivation_letter(category, school, programme, background, prog_info="",
                            full_name="", contact_line=""):
    """Return DOCX bytes for a motivation/scholarship letter grounded in live school research."""
    ctx = retrieve_career_guidance(f"education guidance {category} {programme} Africa scholarship motivation")
    research = web_research(f"{school} {programme} {category} admissions focus, values, and what they look for") if TAVILY_API_KEY else ""
    cat_guidance = {
        "Undergraduate program": "Emphasize academic passion, achievements/grades, why this programme and school, and goals.",
        "PhD / Doctorate position": "Emphasize research fit with the group, prior research/methods/outputs, and long-term goals.",
        "Scholarship": "Emphasize merit and motivation, leadership, need if relevant, and intended impact for Africa.",
    }.get(category, "Emphasize fit, achievements, and goals.")
    prompt = (f"You are an expert admissions/scholarship writing coach.\n\n"
              f"APPLICATION TYPE: {category}\nINSTITUTION: {school}\nPROGRAMME/SCHOLARSHIP: {programme}\n"
              f"APPLICANT NAME: {full_name or '(not provided)'}\nCONTACT: {contact_line or '(not provided)'}\n\n"
              f"APPLICANT BACKGROUND (ground truth):\n{background[:6000]}\n\n"
              f"PROGRAMME DETAILS:\n{prog_info[:4000]}\n\n"
              f"SCHOOL RESEARCH (verified; do not invent beyond this):\n{research or '(none available)'}\n\n"
              f"CATEGORY GUIDANCE: {cat_guidance}\n\n"
              "RESPOND ONLY WITH VALID JSON: {"
              '"full_name": "", "credentials": "", "contact_line": "", '
              f'"date": "{datetime.now().strftime("%B %d, %Y")}", '
              f'"addressee_lines": ["Admissions / Selection Committee", "{school}"], "re_line": "{category}: {programme}", '
              '"salutation": "Dear Members of the Selection Committee,", "body_paragraphs": ['
              '"Opening: what you are applying for + a specific strength of the programme (use research).",'
              '"Most relevant background/achievements mapped to what the programme values.",'
              '"A second dimension (research fit / leadership / impact) with concrete evidence.",'
              '"Why THIS institution/programme (use research) and how it fits your goals.",'
              '"Close: restate motivation, note readiness, thank the committee."], '
              '"closing_line": "Yours sincerely,", "signature_name": "", "signature_title": "", "signature_contact": ""}\n\n'
              "RULES: use ONLY facts from the background/programme info; do NOT invent grades/awards/publications; return ONLY the JSON.")
    return generate_premium_cover_letter_docx(safe_llm_call(prompt, ctx, "English"))


def find_courses(interest, level="Beginner", cost_pref="Free & Paid"):
    """Return a list of verified course recommendations honoring the cost preference."""
    ctx = retrieve_career_guidance(f"skills development training courses {interest} African youth")
    if cost_pref.startswith("Free"):
        guidance = ("Recommend ONLY free-to-access courses; prefer freeCodeCamp, Khan Academy, YouTube, Alison, "
                    "MIT OpenCourseWare, Class Central, and free-to-audit Coursera/edX. Do NOT include Udemy/Udacity/LinkedIn Learning.")
    elif cost_pref.startswith("Paid"):
        guidance = "Recommend paid courses/certifications; prefer Udemy, Udacity, LinkedIn Learning, paid Coursera/edX."
    else:
        guidance = "Include a mix of free and paid options."
    prompt = (f"Recommend 8 real learning resources for someone who wants to learn: {interest}\nLevel: {level}\n\n"
              f"COST REQUIREMENT: {guidance}\n"
              'Free means accessible at no cost (free-to-audit counts as Free).\n\n'
              'Return ONLY a JSON array of objects: {"title": "", "provider": "one of Coursera, edX, Udemy, Udacity, '
              'Class Central, freeCodeCamp, Khan Academy, LinkedIn Learning, YouTube, Alison, FutureLearn, MIT OpenCourseWare", '
              '"cost": "Free|Paid", "level": "", "duration": "", "why": ""}. '
              "Do NOT include URLs (the app builds them). Return ONLY the JSON array.")
    try:
        recs = _extract_json(safe_llm_call(prompt, ctx, "English"))
    except Exception:
        recs = []
    out = []
    if isinstance(recs, list):
        for r in recs:
            if not isinstance(r, dict) or not str(r.get("title", "")).strip():
                continue
            provider = str(r.get("provider", "")).strip()
            cost = classify_cost(provider, str(r.get("cost", "")))
            if not cost_matches(cost, cost_pref):
                continue
            title = str(r["title"]).strip()
            url = provider_search_url(provider, title)
            if not verify_url(url):
                url = provider_search_url("class central", title)
            out.append({"title": title, "provider": provider or "Class Central", "cost": cost,
                        "level": str(r.get("level", "")).strip(), "duration": str(r.get("duration", "")).strip(),
                        "why": str(r.get("why", "")).strip(), "url": url})
    return out[:6]


def find_jobs(role, discipline="", location="", experience="", work_mode="",
              period="", include_ngo=True):
    """Return a list of verified current job openings."""
    yr = datetime.now().year
    parts = [role]
    for extra in (discipline, experience, work_mode):
        if extra and not extra.lower().startswith("any"):
            parts.append(extra)
    parts.append("jobs")
    if location:
        parts.append("in " + location)
    parts.append(f"{yr} apply")
    query = " ".join(parts)
    time_range = {"Past 24 hours": "day", "Past week": "week", "Past month": "month"}.get(period, "")
    results = web_job_search(query, time_range=time_range, domains=_JOB_BOARDS, max_results=8)
    if include_ngo:
        results += web_job_search(query + " NGO OR United Nations OR international organization",
                                  time_range=time_range, domains=_NGO_IO_BOARDS, max_results=8)
    seen, uniq = set(), []
    for r in results:
        u = r.get("url", "")
        if u and u not in seen:
            seen.add(u); uniq.append(r)
    out = []
    for r in uniq:
        if verify_url(r["url"]):
            title, domain = _job_label(r)
            out.append({"title": title, "source": domain, "url": r["url"],
                        "snippet": " ".join((r.get("content") or "").split())[:160]})
    return out


def find_opportunities(opp_type, field, region):
    """Return verified live scholarship / PhD / admissions opportunities."""
    yr = datetime.now().year
    query = f"{opp_type} opportunities {field} {region} {yr} {yr + 1} application requirements deadline"
    return [{"title": " ".join((l.get("title") or "").split()).strip() or l["url"], "url": l["url"]}
            for l in web_search_links(query, max_results=6) if verify_url(l["url"])]
